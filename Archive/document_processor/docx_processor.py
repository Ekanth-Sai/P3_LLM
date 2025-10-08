# document_processor/docx_processor.py

from docx import Document
from docx.shared import Inches
from docx.opc.constants import RELATIONSHIP_TYPE
import os
import io
from typing import List
from PIL import Image

from .base_processor import BaseDocumentProcessor

class DocxProcessor(BaseDocumentProcessor):
    """Processes .docx files."""

    def extract_text(self) -> str:
        doc = Document(self.file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return "\n".join(full_text)

    def extract_images(self) -> List[bytes]:
        doc = Document(self.file_path)
        image_bytes_list = []
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref: # Simplified check
                if rel.reltype == RELATIONSHIP_TYPE.IMAGE:
                    image_part = rel.target_part
                    image_bytes_list.append(image_part.blob)
        return image_bytes_list

# Example Usage (for testing individual components)
if __name__ == "__main__":
    # Create a dummy docx file for testing
    from docx.shared import Inches
    test_doc = Document()
    test_doc.add_heading('Test Docx Document', level=1)
    test_doc.add_paragraph('This is some text in the document.')
    test_doc.add_paragraph('This is another paragraph.')
    
    # Add a dummy image (you'd replace with a real image path)
    try:
        # Create a dummy image
        img = Image.new('RGB', (60, 30), color = 'red')
        img_byte_arr = io.BytesIO()
        img.save(img_byte_arr, format='PNG')
        img_byte_arr = img_byte_arr.getvalue()
        
        with open("temp_image.png", "wb") as f:
            f.write(img_byte_arr)

        test_doc.add_picture("temp_image.png", width=Inches(1.25))
        print("Added dummy image to test_doc.docx")
    except Exception as e:
        print(f"Could not add dummy image ( Pillow or file issue): {e}")

    test_doc.save("data/test_doc.docx")

    processor = DocxProcessor("data/test_doc.docx")
    print("--- DOCX Text ---")
    print(processor.extract_text())
    print("\n--- DOCX Images (Bytes) ---")
    images = processor.extract_images()
    print(f"Found {len(images)} image(s)")
    if images:
        with open("data/extracted_docx_image.png", "wb") as f:
            f.write(images[0])
        print("First image saved to data/extracted_docx_image.png")
    os.remove("data/test_doc.docx")
    if os.path.exists("temp_image.png"):
        os.remove("temp_image.png")